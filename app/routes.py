from datetime import date, datetime, timedelta
from dateutil.relativedelta import relativedelta
import io
import os
from docx import Document
from flask import Response, flash, jsonify, redirect, render_template, request, send_file, send_from_directory, url_for
from flask_login import current_user,login_required, login_user, logout_user
import numpy_financial as npf
import xlsxwriter
from app import app, db, s, socketio, current_active_users
from app.metodos import send_email_amazon
from app.models import Administradores, Archivos, Auditoria, Creditos, EstadosDeuda, ModalidadesPago, Notification, Pagos, Usuarios
from sqlalchemy import and_,or_, case, exists, func, text, extract
from functools import wraps
from werkzeug.utils import secure_filename

import tempfile
from num2words import num2words
from docx.shared import Inches

from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import subprocess
import time
from sqlalchemy.orm import joinedload

# Ruta base del directorio actual
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# Ruta de la carpeta donde se guardarán los PDFs
UPLOAD_FOLDER = os.path.join(BASE_DIR, '..', 'pdfs')
UPLOAD_FOLDER_PAGOS = os.path.join(BASE_DIR, '..', 'images')

# Asegurarse de que la carpeta exista
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(UPLOAD_FOLDER_PAGOS, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}


@app.before_request
def redirect_ip_to_https():
    if request.host.startswith("3.88.129.139"):
        return redirect(request.url.replace("http://3.88.129.139:5000", "https://reagendarcartera.com"), code=301)


def convert_to_pdf(docx_path, output_dir):
    try:
        result = subprocess.run([
            "libreoffice", "--headless", "--convert-to", "pdf",
            "--outdir", output_dir, docx_path
        ], stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, check=True)

        print("Salida estándar:", result.stdout)
        print("Error estándar:", result.stderr)

        pdf_filename = os.path.splitext(os.path.basename(docx_path))[0] + ".pdf"
        pdf_path = os.path.join(output_dir, pdf_filename)

        if not os.path.exists(pdf_path) or os.path.getsize(pdf_path) == 0:
            raise RuntimeError("La conversión a PDF falló o el archivo está vacío.")

        return pdf_path
    except subprocess.CalledProcessError as e:
        print(f"Error al convertir con LibreOffice: {e.stderr}")
        return None
    except Exception as e:
        print(f"Error general: {str(e)}")
        return None



def check_db_connection(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        try:
            # Intentar una consulta trivial usando text
            db.session.execute(text('SELECT 1'))
        except Exception as e:
            app.logger.error(f"Database connection lost: {e}")
            db.session.rollback()
        return f(*args, **kwargs)
    return decorated_function

@app.route("/login", methods=["GET", "POST"])
@check_db_connection
def login():
    if current_user.is_authenticated:
        return redirect(url_for("list_users"))
    
    mostrar_alerta = False

    if request.method == "POST":      
        email = request.form.get("email")
        password = request.form.get("password")
  
        user = Administradores.query.filter_by(correo=email).first()
        if user and user.check_password(password):
            login_user(user)
            # logging.info(f"Inicio de sesión exitoso: usuario {user.correo}")
            return redirect(url_for("list_users"))
            
        else:
            # logging.info(f"Inicio de sesión fallido, credenciales invalidas: usuario {user.correo}")
            mostrar_alerta = True
            return render_template(
                "login.html",
                mostrar_alerta=mostrar_alerta,
            )

    elif request.method == "GET":
        return render_template("login.html")
    
@app.route("/marketing", methods=["GET", "POST"])
def marketing():
    mensaje_exito = False  # Por defecto, no se muestra el mensaje

    if request.method == "POST":
        # Capturar datos del formulario
        nombre = request.form.get("nombre")
        cedula = request.form.get("cedula")
        numero = request.form.get("numero")
        correo = request.form.get("correo")

        # Aquí podrías hacer algo con los datos (guardar en BD, etc.)
        send_email_amazon("financiamientoreagendar@gmail.com", "Interesado en financiamiento", f"Nombre: {nombre} \nCedula: {cedula} \nNúmero: {numero} \nCorreo: {correo}")

        mensaje_exito = True  # Activar el mensaje de éxito

        return render_template("marketing.html", mensaje_exito=mensaje_exito)

    return render_template("marketing.html", mensaje_exito=mensaje_exito)

@app.route('/reset_password', methods=['GET', 'POST'])
@check_db_connection
def reset_password():
    mostrar_alerta = False
    mostrar_exito = False

    if request.method == 'POST':
        email = request.form['email']
        user = Administradores.query.filter_by(correo=email).first()
        if not user:
            mostrar_alerta = True
        else:

            token = s.dumps(email, salt='password-reset-salt')
            link = url_for('reset_with_token', token=token, _external=True)
            
            subject = "Solicitud de restablecimiento de contraseña"
            body_text = f"Su enlace para restablecer su contraseña es {link}"

            if send_email_amazon(email, subject, body_text):
                mostrar_exito = True
                    

    return render_template('reset_password.html',mostrar_exito=mostrar_exito, mostrar_alerta=mostrar_alerta)

@app.route('/reset/<token>', methods=['GET', 'POST'])
@check_db_connection
def reset_with_token(token):
    try:
        email = s.loads(token, salt='password-reset-salt', max_age=3600)
    except:
        flash('The reset link is invalid or has expired', 'danger')
        return redirect(url_for('reset_password'))

    if request.method == 'POST':
        password = request.form['password']
        user = Administradores.query.filter_by(correo=email).first()
        user.set_password(password)
        db.session.commit()

        flash('Your password has been updated!', 'success')
        return redirect(url_for('login'))

    return render_template('reset_with_token.html', token=token)




@app.route("/payment/<int:id>", methods=["GET", "POST"])
@check_db_connection
@login_required
def payment(id):
    
    if current_user.id_administrador != 1 and current_user.id_administrador !=2:
        return "No tienes permiso", 403
    pago = Pagos.query.filter_by(id_pago=id).first()
    
    if not pago:
        return "Pago no encontrado", 404
    
    usuario = (
        db.session.query(Usuarios.id_usuario, Usuarios.nombre_completo, Creditos.id_credito)
        .join(Creditos, Creditos.id_usuario == Usuarios.id_usuario)  # Relación Usuarios -> Creditos
        .join(Pagos, Pagos.id_credito == Creditos.id_credito)  # Relación Creditos -> Pagos
        .filter(Pagos.id_pago == id)  # Filtrar por id_pago
        .first()
    )
    
    nombre_completo = usuario.nombre_completo
    id_usuario = usuario.id_usuario
    id_credito = usuario.id_credito

    variable = ""
    
    mostrar_exito = False
    
    if request.method == "POST":
        fecha_pago_real = request.form.get("fecha_pago_real")
        medio_pago = request.form.get("medio_pago")

        if fecha_pago_real and medio_pago:  # Asegurarse de que se envió una fecha
           

            fecha_pago_programada = pago.fecha_pago_programada
            fecha_pago_real_date = datetime.strptime(fecha_pago_real, "%Y-%m-%d").date()
            dias_mora = max((fecha_pago_real_date - fecha_pago_programada).days, 0)
            # Actualizar en la base de datos
            
            if pago.fecha_pago_real:
                variable = "Actualizar"
            else:
                variable = "Agregar"
                
            pago.fecha_pago_real = fecha_pago_real
            pago.medio = medio_pago
            pago.dias_mora = dias_mora
            pago.id_estado_deuda = 2
            db.session.commit()
            
            # Verificar si es la primera cuota
            if pago.numero_cuota == 1:
                # Obtener la modalidad desde la tabla Creditos
                credito = Creditos.query.filter_by(id_credito=pago.id_credito).first()
                if credito:
                    modalidad = credito.id_modalidad_pago
                    fecha_inicio = fecha_pago_real_date
                    
                    # Actualizar la fecha programada de la cuota 1 a la fecha de inicio
                    pago.fecha_pago_programada = fecha_inicio

                    # Actualizar las fechas programadas para las demás cuotas
                    pagos = Pagos.query.filter(
                        Pagos.id_credito == pago.id_credito,
                        Pagos.numero_cuota > 1
                    ).order_by(Pagos.numero_cuota).all()

                    for i, cuota in enumerate(pagos, start=1):
                        if modalidad == 1:  # Diario
                            cuota.fecha_pago_programada = fecha_inicio + timedelta(days=i)
                        elif modalidad == 2:  # Semanal
                            cuota.fecha_pago_programada = fecha_inicio + timedelta(weeks=i)
                        elif modalidad == 3:  # Quincenal
                            cuota.fecha_pago_programada = fecha_inicio + timedelta(days=15 * i)
                        elif modalidad == 4:  # Mensual
                            cuota.fecha_pago_programada = fecha_inicio + relativedelta(months=i)

                    db.session.commit()
                
             
            cambios = {
                "id credito": id_credito,
                "id pago": id,
                "Fecha de pago": fecha_pago_real,
                "Medio": medio_pago
            }
            
            auditoria = Auditoria(
                id_administrador=current_user.id_administrador,  # El ID del administrador que está registrado en la sesión
                tabla_afectada="Pagos",
                id_usuario= id_usuario,
                cambios=cambios,
                tipo_operacion=variable
            )
            db.session.add(auditoria)
            db.session.commit()
            
            # Verificar si se ha subido un archivo
            if 'file' in request.files:
                file = request.files['file']
                if file.filename != '':
                    # Asegurarse de que el nombre del archivo sea seguro
                    filename = secure_filename(file.filename)
                    # Guardar el archivo con el nombre ID_USUARIO.pdf
                    filepath = os.path.join('images', f"{id}.jpg")
                    file.save(filepath)
            
            
            mostrar_exito = True
        return render_template("payment.html", pago=pago, nombre_completo=nombre_completo,mostrar_exito=mostrar_exito)

        
    elif request.method == "GET":
        return render_template("payment.html", pago=pago,nombre_completo=nombre_completo)
    
    

@app.route("/register_credito/<int:id>", methods=["GET", "POST"])
@check_db_connection
@login_required
def register_credito(id):
    if current_user.id_administrador != 1 and current_user.id_administrador !=2:
        return "No tienes permiso", 403
    mostrar_exito = False
    # Obtener el nombre completo del usuario usando el id
    usuario = Usuarios.query.filter_by(id_usuario=id).first()
    if not usuario:
        return "Usuario no encontrado", 404  # Manejo de error si no existe el usuario

    nombre_completo = usuario.nombre_completo  # Asegúrate de que `nombre_completo` sea un atributo válido
    query = ModalidadesPago.query
    modalidades = query.all()

    if request.method == "POST":
        # Nuevo campo para seleccionar el tipo de crédito
        credit_type = request.form.get("credit_type")  # 'normal' o 'consular'
        
        # Cantidad de adultos y niños
        adults = int(request.form.get("adults", 0))
        children = int(request.form.get("children", 0))
        total_people = adults + children

        # Calcular el costo del servicio
        service_cost = (adults * 1096000) + (children * 694000)

        if credit_type == 'consular':
            # Para crédito consular: se suma (costos consulares - anticipo) por persona = 447500
            amount = service_cost + (total_people * 447500)
        elif credit_type == 'asesoria':
            # Para el nuevo crédito: precios diferentes para adultos y niños
            adult_price = 652000 - 100000  # 552000 por adulto (menos anticipo)
            child_price = 388000 - 100000  # 288000 por niño (menos anticipo)

            amount = (adults * adult_price) + (children * child_price)
        elif credit_type == 'asesoria+consular':
            # Nuevo tipo: mismo precio para adultos y niños, con anticipo restado
            price_after_anticipo = 1432900 - 100000  # 1332900 por persona

            amount = total_people * price_after_anticipo
        else:
            # Para crédito normal: se descuenta 100000 por persona
            amount = service_cost - (total_people * 100000)
            amount = max(0, amount)  # evitar monto negativo

        
        monthly_interest_rate = float(request.form.get("interest_rate").replace(',', '.'))  # Tasa de interés mensual
        installments = int(request.form.get("installments"))  # Número de cuotas
        modalidad = int(request.form.get("modalidad"))  # Modalidad de pago
        date_start = request.form.get("date_start")  # Fecha de inicio
        fecha_date = request.form.get("date_start")

        # Verifica si `date_start` ya es un objeto datetime
        if isinstance(date_start, str):  # Si date_start es una cadena, conviértelo
            date_start = datetime.strptime(date_start, '%Y-%m-%d')  # Suponiendo formato YYYY-MM-DD
        elif isinstance(date_start, datetime):  # Si ya es un objeto datetime, usarlo directamente
            pass
        else:
            return "Formato de fecha no válido", 400  # Si date_start no es ni cadena ni datetime

        try:
            # Cálculo de la cuota fija
            cuota = round(-npf.pmt(monthly_interest_rate / 100, installments, amount))  # Cuota fija calculada con PMT
            if monthly_interest_rate == 0.0:
                total_ganancia = 0
            else:
                total_ganancia = round((cuota * installments) - amount)  # Ganancia total
            # Inicialización de variables para el desglose

            nuevo_credito = Creditos(
                id_usuario=id,
                monto_credito=amount,
                tasa_interes=monthly_interest_rate,
                numero_cuotas=installments,
                id_modalidad_pago=modalidad,
                fecha_credito=date_start,
                ganancia_total=total_ganancia,
                tipo=credit_type
            )
            db.session.add(nuevo_credito)
            db.session.commit()
            db.session.refresh(nuevo_credito)

            saldo = amount

            for i in range(1, installments + 1):
                # Abono a intereses: saldo pendiente * tasa mensual
                abono_intereses = round(saldo * (monthly_interest_rate / 100))  

                # Abono a capital: cuota total - abono a intereses
                abono_capital = round(cuota - abono_intereses)  

                # Actualiza saldo restante
                saldo -= abono_capital  

                # Ajustar el saldo restante en la última cuota para que sea exactamente 0
                if i == installments:
                    abono_capital += saldo  # Asegura que el saldo final sea 0
                    saldo = 0

                # Calcular la fecha programada de pago según la modalidad
                if modalidad == 1:  # Diario
                    fecha_programada = date_start + timedelta(days=i)
                elif modalidad == 2:  # Semanal
                    fecha_programada = date_start + timedelta(weeks=i)
                elif modalidad == 3:  # Quincenal
                    # Suma 15 días para cada incremento
                    fecha_programada = date_start + timedelta(days=15 * i)
                elif modalidad == 4:  # Mensual
                    # Suma i meses de forma exacta
                    fecha_programada = date_start + relativedelta(months=i)

                # Registrar el pago en la base de datos
                pagos = Pagos(
                    id_credito=nuevo_credito.id_credito,
                    numero_cuota=i,
                    fecha_pago_programada=fecha_programada,
                    monto_cuota=cuota,
                    abono_capital=abono_capital,
                    abono_intereses=abono_intereses,
                    saldo_restante=saldo,
                    id_estado_deuda=1  # Estado de deuda "Pagado"
                )
                db.session.add(pagos)

            modalidad_seleccionada = ModalidadesPago.query.get(modalidad)
            nombre_modalidad = modalidad_seleccionada.nombre_modalidad 

            # Registrar el cambio en la tabla Auditoria
            cambios = {
                "id credito": nuevo_credito.id_credito,
                "Nombre completo": nombre_completo,
                "Monto": amount,
                "Tasa de interes": monthly_interest_rate,
                "Número de cuotas": installments,
                "Modalidad": nombre_modalidad,
                "Fecha de credito": fecha_date,
                "Tipo": credit_type
            }

            auditoria = Auditoria(
                id_administrador=current_user.id_administrador,  # El ID del administrador que está registrado en la sesión
                tabla_afectada="Creditos",
                id_usuario=id,
                cambios=cambios,
                tipo_operacion="Agregar"
            )
            db.session.add(auditoria)
            db.session.commit()
            mostrar_exito = True
            
            # Si el crédito es de tipo consular, enviar notificación al admin (ID 1)
            if credit_type == 'consular':
                message = f"Nuevo crédito consular registrado para {nombre_completo}."
                target_user_id = 3  # Siempre se envía al administrador con ID 3
                send_notification(target_user_id, message, credit_id=nuevo_credito.id_credito)
            else:
                print("chao")
        except Exception as error:
            print("Error al insertar datos en la tabla usuarios:", error)

        return render_template(
            "register_credito.html",
            mostrar_exito=mostrar_exito,
            modalidades=modalidades, 
            id_usuario=id, 
            nombre_completo=nombre_completo
        )
    else:  # Si es un GET
        return render_template("register_credito.html", modalidades=modalidades, id_usuario=id, nombre_completo=nombre_completo)


@app.route("/update_credito/<int:id>", methods=["GET", "POST"])
@check_db_connection
@login_required
def update_credito(id):
    if current_user.id_administrador != 1 and current_user.id_administrador !=2:
        return "No tienes permiso", 403
    mostrar_exito = False
    mostrar_alerta = False
    credito = Creditos.query.filter_by(id_credito=id).first()
    if not credito:
        return "Credito no encontrado", 404
    
    id_usuario = credito.id_usuario
    usuario = Usuarios.query.filter_by(id_usuario=id_usuario).first()
    nombre_usuario = usuario.nombre_completo

    if request.method == "POST":
        # Obtener los nuevos valores del formulario
        amount = int(request.form.get("amount").replace('.', ''))
        monthly_interest_rate = float(request.form.get("interest_rate").replace(',', '.'))
        installments = int(request.form.get("installments"))
        modalidad = int(request.form.get("modalidad"))
        date_start = request.form.get("date_start")
        fecha_date = request.form.get("date_start")

        # Verifica si la fecha de inicio es válida
        if isinstance(date_start, str):
            date_start = datetime.strptime(date_start, '%Y-%m-%d')
        elif isinstance(date_start, datetime):
            pass
        else:
            return "Formato de fecha no válido", 400

        try:
            # Calcular valores del crédito
            cuota = round(-npf.pmt(monthly_interest_rate / 100, installments, amount))
            total_ganancia = round((cuota * installments) - amount)

            # Actualizar el crédito
            credito.monto_credito = amount
            credito.tasa_interes = monthly_interest_rate
            credito.numero_cuotas = installments
            credito.id_modalidad_pago = modalidad
            credito.fecha_credito = date_start
            credito.ganancia_total = total_ganancia
            db.session.commit()

            # Preparar variables para cálculos de pagos
            saldo = amount
            # Obtenemos TODOS los pagos, incluso los deshabilitados
            pagos_actuales = Pagos.query.filter_by(id_credito=credito.id_credito).order_by(Pagos.numero_cuota.asc()).all()
            fecha_inicio = date_start

            # Función auxiliar para calcular fecha programada
            def calcular_fecha_programada(i):
                if modalidad == 1:  # Diario
                    return fecha_inicio + timedelta(days=i)
                elif modalidad == 2:  # Semanal
                    return fecha_inicio + timedelta(weeks=i)
                elif modalidad == 3:  # Quincenal
                    return fecha_inicio + timedelta(weeks=2 * i)
                else:  # Mensual
                    return fecha_inicio + timedelta(weeks=4 * i)

           
            def calcular_detalles_pago(saldo, es_ultima_cuota=False):
                abono_intereses = round(saldo * (monthly_interest_rate / 100))
                abono_capital = round(cuota - abono_intereses)
                nuevo_saldo = saldo - abono_capital

                if es_ultima_cuota:
                    # Ajustamos el abono a capital en la última cuota para que el saldo sea exactamente 0
                    abono_capital = saldo
                    nuevo_saldo = 0

                return abono_intereses, abono_capital, nuevo_saldo

            # Función auxiliar para actualizar un pago existente o crear uno nuevo
            def actualizar_o_crear_pago(numero_cuota, saldo, es_ultima_cuota):
                # Buscamos primero si existe un pago para esta cuota (habilitado o no)
                pago = next((p for p in pagos_actuales if p.numero_cuota == numero_cuota), None)
                abono_intereses, abono_capital, nuevo_saldo = calcular_detalles_pago(saldo, es_ultima_cuota)
                
                if not pago:
                    pago = Pagos(
                        id_credito=credito.id_credito,
                        numero_cuota=numero_cuota,
                        id_estado_deuda=1,
                        habilitado=True
                    )
                    db.session.add(pago)
                else:
                    # Si el pago existe, lo habilitamos
                    pago.habilitado = True
                
                pago.fecha_pago_programada = calcular_fecha_programada(numero_cuota)
                pago.monto_cuota = cuota
                pago.abono_capital = abono_capital
                pago.abono_intereses = abono_intereses
                pago.saldo_restante = nuevo_saldo
                
                return nuevo_saldo

            # CASO 1: El número de cuotas aumentó
            if installments > len([p for p in pagos_actuales if p.habilitado]):
                # Primero recalculamos todas las cuotas existentes y habilitadas
                for i in range(1, installments + 1):
                    saldo = actualizar_o_crear_pago(i, saldo, i == installments)

            # CASO 2: El número de cuotas disminuyó
            elif installments < len([p for p in pagos_actuales if p.habilitado]):
                # Deshabilitamos las cuotas que sobran
                for pago in pagos_actuales:
                    if pago.numero_cuota > installments:
                        pago.habilitado = False
                
                # Recalculamos las cuotas que quedan
                for i in range(1, installments + 1):
                    saldo = actualizar_o_crear_pago(i, saldo, i == installments)

            # CASO 3: El número de cuotas es igual
            else:
                # Recalculamos todas las cuotas
                for i in range(1, installments + 1):
                    saldo = actualizar_o_crear_pago(i, saldo, i == installments)

            # Registrar auditoría
            cambios = {
                "id credito": credito.id_credito,
                "Nombre completo": nombre_usuario,
                "Monto": amount,
                "Tasa de interes": monthly_interest_rate,
                "Número de cuotas": installments,
                "Modalidad": ModalidadesPago.query.get(modalidad).nombre_modalidad,
                "Fecha de credito": fecha_date
            }

            auditoria = Auditoria(
                id_administrador=current_user.id_administrador,
                tabla_afectada="Creditos",
                id_usuario=id_usuario,
                cambios=cambios,
                tipo_operacion="Actualizar"
            )
            db.session.add(auditoria)
            db.session.commit()

            mostrar_exito = True

        except Exception as error:
            print("Error al actualizar el crédito:", error)
            db.session.rollback()

        return render_template(
            "update_credito.html",
            nombre_usuario=nombre_usuario,
            credito=credito,
            mostrar_alerta=mostrar_alerta,
            mostrar_exito=mostrar_exito,
        )

    return render_template(
        "update_credito.html",
        nombre_usuario=nombre_usuario,
        credito=credito,
        mostrar_alerta=mostrar_alerta,
        mostrar_exito=mostrar_exito,
    )
        

@app.route("/register_user", methods=["GET", "POST"])
@check_db_connection
@login_required
def register_user():
    if current_user.id_administrador != 1 and current_user.id_administrador !=2:
        return "No tienes permiso", 403
    mostrar_alerta = False
    mostrar_exito = False
    mostrar_alerta2 = False

    if request.method == "POST":
        cedula = request.form.get("cedula")
        correo = request.form.get("correo")

        usuario_existente = Usuarios.query.filter((Usuarios.cedula == cedula) | (Usuarios.correo == correo)).first() 
        if usuario_existente: 
            mostrar_alerta = usuario_existente.cedula == cedula 
            mostrar_alerta2 = usuario_existente.correo == correo
        else:
            try:

                nuevo_usuario = Usuarios(
                    nombre_completo=request.form.get("nombre_completo"),
                    cedula=cedula,
                    direccion=request.form.get("direccion"),
                    celular=request.form.get("celular"),
                    correo=request.form.get("correo"),
                    expedicion=request.form.get("expedicion"),
                    residencia=request.form.get("residencia")
                )
                db.session.add(nuevo_usuario)
                db.session.commit()
                db.session.refresh(nuevo_usuario) 
                
                # Verificar si se ha subido un archivo
                if 'file' in request.files:
                    file = request.files['file']
                    if file.filename != '':
                        # Asegurarse de que el nombre del archivo sea seguro
                        filename = secure_filename(file.filename)
                        # Guardar el archivo con el nombre ID_USUARIO.pdf
                        filepath = os.path.join('pdfs', f"{nuevo_usuario.id_usuario}.pdf")
                        file.save(filepath)
                        
                # Registrar el cambio en la tabla Auditoria
                cambios = {
                    "Nombre completo": nuevo_usuario.nombre_completo,
                    "Cedula": nuevo_usuario.cedula,
                    "Direccion": nuevo_usuario.direccion,
                    "Celular": nuevo_usuario.celular,
                    "Correo": nuevo_usuario.correo
                }
                
                auditoria = Auditoria(
                    id_administrador=current_user.id_administrador,  # El ID del administrador que está registrado en la sesión
                    tabla_afectada="Usuarios",
                    id_usuario=nuevo_usuario.id_usuario,
                    cambios=cambios,
                    tipo_operacion="Agregar"
                )
                db.session.add(auditoria)
                db.session.commit()
                mostrar_exito = True
                
            except Exception as error:
                print("Error al insertar datos en la tabla usuarios:", error)

    
        # logging.info(f"Resultado del registro por excel: usuario {current_user.correo}")
        return render_template(
            "register_user.html",
            mostrar_alerta=mostrar_alerta,
            mostrar_exito=mostrar_exito,
            mostrar_alerta2=mostrar_alerta2
        )
    elif request.method == "GET":
        # logging.info(f"Ingreso al registro por excel: usuario {current_user.correo}")
        return render_template("register_user.html")



@app.route("/register_employee", methods=["GET", "POST"])
@check_db_connection
@login_required
def register_employee():
    if current_user.id_administrador !=2:
        return "No tienes permiso", 403
        
        
    mostrar_alerta = False
    mostrar_exito = False

    if request.method == "POST":
        correo = request.form.get("email")

        usuario_existente = Administradores.query.filter(Administradores.correo == correo).first() 
        if usuario_existente: 
            mostrar_alerta = usuario_existente.correo == correo 
        else:
            try:

                nuevo_empleado = Administradores(
                    correo=request.form.get("email"),
                    id_rol=1
                )
                nuevo_empleado.set_password(request.form.get("password"))
                db.session.add(nuevo_empleado)
                db.session.commit()                
                mostrar_exito = True
                
            except Exception as error:
                print("Error al insertar datos en la tabla usuarios:", error)

    
        # logging.info(f"Resultado del registro por excel: usuario {current_user.correo}")
        return render_template(
            "register_employee.html",
            mostrar_alerta=mostrar_alerta,
            mostrar_exito=mostrar_exito,
        )
    elif request.method == "GET":
        # logging.info(f"Ingreso al registro por excel: usuario {current_user.correo}")
        return render_template("register_employee.html")

@app.route("/logout")
@check_db_connection
@login_required
def logout():
    logout_user()
    flash("Has cerrado sesión", "success")
    return redirect(url_for("login"))


@app.route("/list_users", methods=["GET"])
@check_db_connection
@login_required
def list_users():
    filtro = request.args.get("filtro", "")
    mes_filtro = request.args.get("mes_filtro", "")
    anio_filtro = request.args.get("anio_filtro", "")
    pagina_actual = request.args.get("pagina", default=1, type=int)
    usuarios_por_pagina = 10

    try:
        # Query base - solo usuarios habilitados
        query = Usuarios.query.filter_by(habilitado=1)
        
        # Filtro por texto (nombre, cedula, celular)
        if filtro:
            query = query.filter(
                or_(
                    Usuarios.nombre_completo.ilike(f"%{filtro}%"),
                    Usuarios.cedula.ilike(f"%{filtro}%"),
                    Usuarios.celular.ilike(f"%{filtro}%")
                )
            )
        
        # Filtro por mes
        if mes_filtro:
            query = query.filter(extract('month', Usuarios.fecha_creacion) == int(mes_filtro))
        
        # Filtro por año
        if anio_filtro:
            query = query.filter(extract('year', Usuarios.fecha_creacion) == int(anio_filtro))
        
        # Obtener años disponibles para el dropdown
        anios_disponibles = db.session.query(
            extract('year', Usuarios.fecha_creacion).label('anio')
        ).filter(
            Usuarios.habilitado == 1,
            Usuarios.fecha_creacion.isnot(None)
        ).distinct().order_by(
            extract('year', Usuarios.fecha_creacion).desc()
        ).all()
        
        anios_disponibles = [anio.anio for anio in anios_disponibles if anio.anio]
        
        # Paginación
        total_usuarios = query.count()
        total_paginas = (total_usuarios - 1) // usuarios_por_pagina + 1 if total_usuarios > 0 else 1
        inicio = (pagina_actual - 1) * usuarios_por_pagina
        users_paginados = query.offset(inicio).limit(usuarios_por_pagina).all()
        
        # Páginas visibles para la paginación
        paginas_visibles = []
        for pagina in range(
            max(1, pagina_actual - 2), min(total_paginas + 1, pagina_actual + 3)
        ):
            paginas_visibles.append(pagina)

        return render_template(
            "list_users.html",
            data=users_paginados,
            pagina_actual=pagina_actual,
            total_usuarios=total_usuarios,
            total_paginas=total_paginas,
            usuarios_por_pagina=usuarios_por_pagina,
            filtro=filtro,
            mes_filtro=mes_filtro,
            anio_filtro=anio_filtro,
            anios_disponibles=anios_disponibles,
            paginas_visibles=paginas_visibles,
        ), 200

    except Exception as e:
        print("Error al obtener los usuarios:", e)
        # En caso de error, mostrar una página en blanco o una página de error
        return render_template(
            "list_users.html",
            data=[],
            pagina_actual=pagina_actual,
            total_usuarios=0,
            total_paginas=0,
            usuarios_por_pagina=usuarios_por_pagina,
            filtro=filtro,
            mes_filtro="",
            anio_filtro="",
            anios_disponibles=[],
            paginas_visibles=[],
        ), 200
    
# Ruta Flask corregida
@app.route('/send_multiple_emails', methods=['POST'])
@check_db_connection
@login_required
def send_multiple_emails():
    if current_user.id_administrador != 1 and current_user.id_administrador !=2:
        return "No tienes permiso", 403
    try:
        # Eliminamos la lectura duplicada de request.get_json()
        data = request.get_json()
        
        if not data or not isinstance(data, dict):
            return jsonify({"success": False, "message": "Datos inválidos"}), 400

        ids_pago = data.get("ids", [])
        
        if not ids_pago or not isinstance(ids_pago, list):
            return jsonify({"success": False, "message": "No se proporcionaron IDs válidos"}), 400

        resultados = []

        for id_pago in ids_pago:
            try:
                pago = db.session.query(
                    Pagos,
                    Usuarios
                ).join(
                    Creditos, Pagos.id_credito == Creditos.id_credito
                ).join(
                    Usuarios, Creditos.id_usuario == Usuarios.id_usuario
                ).filter(
                    Pagos.id_pago == id_pago
                ).first()

                if not pago:
                    resultados.append({"id_pago": id_pago, "status": "Pago no encontrado"})
                    continue

                pago_data = pago[0]
                usuario_data = pago[1]

                # Validación adicional de datos
                if not usuario_data.correo:
                    resultados.append({"id_pago": id_pago, "status": "Correo no disponible"})
                    continue

                email = usuario_data.correo
                subject = "Notificación de mora en su pago"
                body_text = (
                    f"Estimado(a) {usuario_data.nombre_completo},\n\n"
                    f"Le informamos que se encuentra en mora con el pago correspondiente a la cuota #{pago_data.numero_cuota}. "
                    f"Fecha programada para el pago: {pago_data.fecha_pago_programada.strftime('%d/%m/%Y')}.\n"
                    f"Días en mora: {pago_data.dias_mora}.\n\n"
                    f"Por favor, regularice su situación a la brevedad para evitar inconvenientes mayores.\n\n"
                    f"Atentamente,\nReagendar."
                )

                enviado = send_email_amazon(email, subject, body_text)
                resultados.append({
                    "id_pago": id_pago, 
                    "status": "Enviado" if enviado else "Error al enviar",
                    "email": email
                })

            except Exception as e:
                resultados.append({
                    "id_pago": id_pago,
                    "status": f"Error: {str(e)}"
                })

        return jsonify({
            "success": True,
            "resultados": resultados,
            "total_enviados": len([r for r in resultados if r["status"] == "Enviado"])
        }), 200

    except Exception as e:
        return jsonify({
            "success": False, 
            "message": f"Error interno del servidor: {str(e)}"
        }), 500



@app.route("/list_mora", methods=["GET"])
@check_db_connection
@login_required
def list_mora():
    if current_user.id_administrador != 1 and current_user.id_administrador !=2:
        return "No tienes permiso", 403
    pagina_actual = request.args.get("pagina", default=1, type=int)
    usuarios_por_pagina = 10

    try:
        # Consulta para obtener todos los pagos en mora
        query = db.session.query(
            Pagos.id_pago,
            Usuarios.nombre_completo,
            Pagos.numero_cuota,
            Pagos.fecha_pago_programada,
            Pagos.fecha_pago_real,
            Pagos.monto_cuota,
            Pagos.abono_capital,
            Pagos.abono_intereses,
            Pagos.saldo_restante,
            EstadosDeuda.nombre_estado,
            Usuarios.id_usuario,
            Creditos.id_credito,
            Pagos.dias_mora
        ).select_from(Pagos) \
        .join(Creditos, Pagos.id_credito == Creditos.id_credito) \
        .join(Usuarios, Creditos.id_usuario == Usuarios.id_usuario) \
        .join(EstadosDeuda, Pagos.id_estado_deuda == EstadosDeuda.id_estado) \
        .filter(Pagos.id_estado_deuda == 3) \
        .filter(Pagos.dias_mora > 0) \
        .order_by(Pagos.dias_mora.desc())

        # Paginación
        total_usuarios = query.count()
        total_paginas = (total_usuarios - 1) // usuarios_por_pagina + 1
        inicio = (pagina_actual - 1) * usuarios_por_pagina
        users_paginados = query.offset(inicio).limit(usuarios_por_pagina).all()

        # Páginas visibles para la navegación
        paginas_visibles = [
            pagina for pagina in range(
                max(1, pagina_actual - 2), min(total_paginas + 1, pagina_actual + 3)
            )
        ]

        # Renderizamos la plantilla
        return render_template(
            "list_mora.html",
            data=users_paginados,
            pagina_actual=pagina_actual,
            total_usuarios=total_usuarios,
            total_paginas=total_paginas,
            usuarios_por_pagina=usuarios_por_pagina,
            paginas_visibles=paginas_visibles,
        ), 200

    except Exception as e:
        print("Error al obtener los pagos en mora:", e)
        return render_template(
            "list_mora.html",
            data=[],
            pagina_actual=pagina_actual,
            total_usuarios=0,
            total_paginas=0,
            usuarios_por_pagina=usuarios_por_pagina,
            paginas_visibles=[],
        ), 200


@app.route("/list_pagos/<int:id>", methods=["GET"])
@check_db_connection
@login_required
def list_pagos(id):
    # filtro = request.args.get("filtro", "")
    pagina_actual = request.args.get("pagina", default=1, type=int)
    usuarios_por_pagina = 10

    try:
        
        query = db.session.query(
            Pagos.id_pago,
            Usuarios.nombre_completo,
            Pagos.numero_cuota,
            Pagos.fecha_pago_programada,
            Pagos.fecha_pago_real,
            Pagos.monto_cuota,
            Pagos.abono_capital,
            Pagos.abono_intereses,
            Pagos.saldo_restante,
            EstadosDeuda.nombre_estado,
            Usuarios.id_usuario,
            Creditos.id_credito,
            Pagos.dias_mora
        ).select_from(Pagos) \
        .join(Creditos, Pagos.id_credito == Creditos.id_credito) \
        .join(Usuarios, Creditos.id_usuario == Usuarios.id_usuario) \
        .join(EstadosDeuda, Pagos.id_estado_deuda == EstadosDeuda.id_estado) \
        .filter(Creditos.id_credito == id) \
        .filter(Pagos.habilitado == True) \
        .order_by(Pagos.id_pago.asc())


        total_usuarios = query.count()
        total_paginas = (total_usuarios - 1) // usuarios_por_pagina + 1
        inicio = (pagina_actual - 1) * usuarios_por_pagina
        users_paginados = query.offset(inicio).limit(usuarios_por_pagina).all()
        results = query.all()


        paginas_visibles = []
        for pagina in range(
            max(1, pagina_actual - 2), min(total_paginas + 1, pagina_actual + 3)
        ):
            paginas_visibles.append(pagina)

        return render_template(
            "list_pagos.html",
            data=users_paginados,
            pagina_actual=pagina_actual,
            total_usuarios=total_usuarios,
            total_paginas=total_paginas,
            usuarios_por_pagina=usuarios_por_pagina,
            # filtro=filtro,
            paginas_visibles=paginas_visibles,
        ), 200

    except Exception as e:
        print("Error al obtener los usuarios:", e)
        # En caso de error, mostrar una página en blanco o una página de error
        return render_template(
            "list_pagos.html",
            data=[],
            pagina_actual=pagina_actual,
            total_usuarios=0,
            total_paginas=0,
            usuarios_por_pagina=usuarios_por_pagina,
            # filtro=filtro,
            paginas_visibles=[],
        ), 200


@app.route('/download-pdf/<int:id>', methods=['POST'])
def download_pdf(id):
    if current_user.id_administrador != 1 and current_user.id_administrador !=2:
        return "No tienes permiso", 403
    
    try:
        credito = Creditos.query.get_or_404(id)
        usuario = credito.usuario

        def numero_a_texto(numero):
            return num2words(numero, lang='es').upper()

        # Formatear valores monetarios
        monto_credito = f"{credito.monto_credito:,.0f}".replace(",", ".")
        valor_cuota = int(credito.monto_credito / credito.numero_cuotas)
        valor_cuota_formato = f"{valor_cuota:,.0f}".replace(",", ".")
        
        # Formatear la fecha
        fecha = credito.fecha_credito
        dia_texto = num2words(fecha.day, lang='es').capitalize()
        meses = ["enero", "febrero", "marzo", "abril", "mayo", "junio", 
                 "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
        mes_texto = meses[fecha.month - 1]

        # Preparar los datos para el documento
        datos = {
            "{{NOMBRE}}": usuario.nombre_completo.upper(),
            "{{CEDULA}}": usuario.cedula,
            "{{RESIDENCIA}}": usuario.residencia.upper(),
            "{{DIRECCION}}": usuario.direccion.upper(),
            "{{MONTO_TOTAL}}": monto_credito,
            "{{ANTICIPO}}": "100.000",
            "{{NUMERO_CUOTAS}}": str(credito.numero_cuotas),
            "{{NUMERO_CUOTAS_TEXTO}}": num2words(credito.numero_cuotas, lang='es'),
            "{{VALOR_CUOTA}}": valor_cuota_formato,
            "{{MONTO_TOTAL_TEXTO}}": numero_a_texto(credito.monto_credito),
            "{{ANTICIPO_TEXTO}}": "CIEN MIL",
            "{{MONTO_COSTOS_CONSULARES_TEXTO}}": "OCHOCIENTOS TREINTA Y DOS MIL QUINIENTOS",
            "{{MONTO_COSTOS_CONSULARES}}": "832.500",
            "{{MONTO_VALOR_TEXTO}}": numero_a_texto(credito.monto_credito),
            "{{MONTO_VALOR}}": monto_credito,
            "{{NUMERO_CREDITO}}": str(credito.id_credito),
            "{{FECHA_FIRMA}}": fecha.strftime("%d/%m/%Y"),
            "{{FECHA_FIRMA_DIA_TEXTO}}": dia_texto,
            "{{FECHA_FIRMA_DIA}}": str(fecha.day),
            "{{FECHA_FIRMA_MES_TEXTO}}": mes_texto,
            "{{FECHA_FIRMA_AÑO}}": str(fecha.year),
            "{{EXPEDICION}}": usuario.expedicion.upper(),
            "{{CORREO}}": usuario.correo.lower()
        }

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as docx_temp:
            
            ruta_docx = os.path.join(os.getcwd(), "app", "CONTRATO_DE_PRESTACION_DE_SERVICIOS_DE_FINANCIAMIENTO.docx")

            if os.path.exists(ruta_docx):
                contrato = Document(ruta_docx)
                print("Documento cargado correctamente.")
            else:
                print("El archivo no existe en la ruta:", ruta_docx)



            def insertar_imagen(parrafo, imagen_path):
                """
                Inserta una imagen manteniendo la alineación original del párrafo
                """
                try:
                    # Guardar la alineación original antes de limpiar
                    alineacion_original = parrafo.alignment
                    
                    # Limpiar el contenido actual del párrafo
                    for run in parrafo.runs:
                        run.text = ""
                    
                    # Agregar la imagen con tamaño controlado
                    run = parrafo.add_run()
                    run.add_picture(imagen_path, width=Inches(2))
                    
                    # Restaurar la alineación original
                    parrafo.alignment = alineacion_original
                    
                    return True
                except Exception as e:
                    print(f"Error al insertar imagen: {str(e)}")
                    return False

            def get_alineacion_texto(texto):
                """
                Determina la alineación basada en el marcador
                """
                if "{{IMAGEN_DERECHA}}" in texto:
                    return WD_ALIGN_PARAGRAPH.RIGHT
                elif "{{IMAGEN_IZQUIERDA}}" in texto:
                    return WD_ALIGN_PARAGRAPH.LEFT
                return None

            def reemplazar_en_parrafo(parrafo):
                """
                Reemplaza el texto y maneja la inserción de imágenes respetando alineación
                """
                texto_completo = "".join(run.text for run in parrafo.runs)
                
                # Verificar todos los posibles marcadores de imagen
                marcadores_imagen = ["{{IMAGEN_DERECHA}}", "{{IMAGEN_IZQUIERDA}}"]
                
                for marcador in marcadores_imagen:
                    if marcador in texto_completo:
                        # Si se especifica una alineación especial en el marcador
                        alineacion = get_alineacion_texto(texto_completo)
                        if alineacion is not None:
                            parrafo.alignment = alineacion
                        # Si no hay alineación especial, se mantiene la del párrafo
                        
                        return insertar_imagen(parrafo, os.path.join(os.getcwd(), "app", "firma.jpg"))

                # Proceder con el reemplazo normal de texto
                for placeholder, valor in datos.items():
                    texto_completo = texto_completo.replace(placeholder, valor)
                
                # Si el texto no ha cambiado, no hacemos nada
                if texto_completo == "".join(run.text for run in parrafo.runs):
                    return False
                
                # Limpiar el párrafo actual
                for run in parrafo.runs:
                    run.text = ""
                
                # Agregar el nuevo texto con el formato original
                if texto_completo:
                    run = parrafo.add_run(texto_completo)
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
                
                return True

            # Procesar todos los párrafos del documento
            for parrafo in contrato.paragraphs:
                reemplazar_en_parrafo(parrafo)

            # Procesar todas las tablas del documento
            for tabla in contrato.tables:
                for fila in tabla.rows:
                    for celda in fila.cells:
                        for parrafo in celda.paragraphs:
                            reemplazar_en_parrafo(parrafo)

            # Guardar el documento modificado
            contrato.save(docx_temp.name)

        pdf_temp = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
        pdf_temp.close()
        pdf_path = convert_to_pdf(docx_temp.name, os.path.dirname(pdf_temp.name))
        if not pdf_path:
            raise Exception("Error en la conversión de DOCX a PDF")

        with open(pdf_path, 'rb') as pdf_file:
            pdf_content = pdf_file.read()

        os.unlink(docx_temp.name)
        os.unlink(pdf_path)

        buffer = io.BytesIO(pdf_content)
        buffer.seek(0)

        return send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f'Contrato_{id}.pdf'
        )


    except Exception as e:
        return f"Error al generar el PDF: {str(e)}", 500
    


@app.route("/list_credito/<int:id>", methods=["GET"])
@check_db_connection
@login_required
def list_credito(id):
    pagina_actual = request.args.get("pagina", default=1, type=int)
    usuarios_por_pagina = 10

    try:
        # 1) Consulta base con todos los campos necesarios
        query_sql = db.session.query(
            Creditos.id_credito,
            Usuarios.nombre_completo,
            Creditos.monto_credito,
            Creditos.tasa_interes,
            Creditos.numero_cuotas,
            ModalidadesPago.nombre_modalidad,
            Creditos.fecha_credito,
            Creditos.fecha_creacion,
            Creditos.fecha_actualizacion,
            Creditos.tipo,  # Añadimos el tipo para poder calcular numero_personas
            # Conteo de cuotas pagadas
            func.count(
                case(
                    (Pagos.id_estado_deuda == 2, 1),
                    else_=None
                )
            ).label('cuotas_pagadas'),
            # Flag de existencia de docs subidos por admin=3
            exists().where(
                and_(
                    Archivos.id_usuario == Usuarios.id_usuario,
                    Archivos.id_subidor == 3
                )
            ).label('tiene_docs_id3')
        ).join(
            Usuarios, Creditos.id_usuario == Usuarios.id_usuario
        ).join(
            ModalidadesPago, Creditos.id_modalidad_pago == ModalidadesPago.id_modalidad
        ).outerjoin(
            Pagos,
            (Creditos.id_credito == Pagos.id_credito) & (Pagos.habilitado == True)
        ).filter(
            Usuarios.id_usuario == id,
            Creditos.habilitado == True
        ).group_by(
            Creditos.id_credito,
            Usuarios.nombre_completo,
            Creditos.monto_credito,
            Creditos.tasa_interes,
            Creditos.numero_cuotas,
            ModalidadesPago.nombre_modalidad,
            Creditos.fecha_credito,
            Creditos.fecha_creacion,
            Creditos.fecha_actualizacion,
            Creditos.tipo,
            Usuarios.id_usuario,
        ).order_by(
            Creditos.id_credito.asc()
        )

        # 2) Paginación
        total_usuarios = query_sql.count()
        total_paginas = (total_usuarios - 1) // usuarios_por_pagina + 1
        inicio = (pagina_actual - 1) * usuarios_por_pagina

        users_raw = query_sql.offset(inicio).limit(usuarios_por_pagina).all()
        
        # 3) Convertir a diccionarios y añadir numero_personas
        users_paginados = []
        for user in users_raw:
            # Convertir a diccionario para poder modificar
            user_dict = user._asdict() if hasattr(user, '_asdict') else {
                col: getattr(user, col) for col in user._fields
            } if hasattr(user, '_fields') else dict(user._mapping)
            
            # Calcular numero_personas basado en tipo y monto_credito
            if user.tipo == 'normal':
                precio_adulto = 1096000 - 100000  # 996,000
                precio_nino = 694000 - 100000     # 594,000
                
                # Intentar encontrar la combinación exacta de adultos y niños
                encontrado = False
                
                # Probar diferentes combinaciones (hasta 5 personas total)
                for num_adultos in range(6):
                    for num_ninos in range(6 - num_adultos):
                        if num_adultos + num_ninos > 0:  # Al menos una persona
                            total_calculado = (num_adultos * precio_adulto) + (num_ninos * precio_nino)
                            if total_calculado == user.monto_credito:
                                user_dict['numero_personas'] = num_adultos + num_ninos
                                encontrado = True
                                break
                    if encontrado:
                        break
                
                # Si no encontramos combinación exacta, aproximar
                if not encontrado:
                    # Primero intentar solo adultos
                    if user.monto_credito % precio_adulto == 0:
                        user_dict['numero_personas'] = user.monto_credito // precio_adulto
                    # Luego intentar solo niños
                    elif user.monto_credito % precio_nino == 0:
                        user_dict['numero_personas'] = user.monto_credito // precio_nino
                    else:
                        # Si no es divisible, buscar la mejor aproximación
                        mejor_diferencia = float('inf')
                        mejor_personas = 1
                        
                        for num_personas in range(1, 6):  # Máximo 5 personas
                            # Probar todas las combinaciones para este número de personas
                            for num_adultos in range(num_personas + 1):
                                num_ninos = num_personas - num_adultos
                                total = (num_adultos * precio_adulto) + (num_ninos * precio_nino)
                                diferencia = abs(total - user.monto_credito)
                                
                                if diferencia < mejor_diferencia:
                                    mejor_diferencia = diferencia
                                    mejor_personas = num_personas
                        
                        user_dict['numero_personas'] = mejor_personas

            elif user.tipo == 'consular':
                precio_adulto = 1928500 - 385000  # 1,543,500
                precio_nino = 1526500 - 385000    # 1,141,500
                
                # Mismo enfoque: buscar combinación exacta primero
                encontrado = False
                
                for num_adultos in range(6):
                    for num_ninos in range(6 - num_adultos):
                        if num_adultos + num_ninos > 0:
                            total_calculado = (num_adultos * precio_adulto) + (num_ninos * precio_nino)
                            if total_calculado == user.monto_credito:
                                user_dict['numero_personas'] = num_adultos + num_ninos
                                encontrado = True
                                break
                    if encontrado:
                        break
                
                if not encontrado:
                    if user.monto_credito % precio_adulto == 0:
                        user_dict['numero_personas'] = user.monto_credito // precio_adulto
                    elif user.monto_credito % precio_nino == 0:
                        user_dict['numero_personas'] = user.monto_credito // precio_nino
                    else:
                        mejor_diferencia = float('inf')
                        mejor_personas = 1
                        
                        for num_personas in range(1, 6):  # Máximo 5 personas
                            for num_adultos in range(num_personas + 1):
                                num_ninos = num_personas - num_adultos
                                total = (num_adultos * precio_adulto) + (num_ninos * precio_nino)
                                diferencia = abs(total - user.monto_credito)
                                
                                if diferencia < mejor_diferencia:
                                    mejor_diferencia = diferencia
                                    mejor_personas = num_personas
                        
                        user_dict['numero_personas'] = mejor_personas

            elif user.tipo == 'asesoria':
                precio_adulto = 652000 - 100000  # 552,000
                precio_nino = 388000 - 100000    # 288,000
                
                # Mismo enfoque
                encontrado = False
                
                for num_adultos in range(21):
                    for num_ninos in range(21 - num_adultos):
                        if num_adultos + num_ninos > 0:
                            total_calculado = (num_adultos * precio_adulto) + (num_ninos * precio_nino)
                            if total_calculado == user.monto_credito:
                                user_dict['numero_personas'] = num_adultos + num_ninos
                                encontrado = True
                                break
                    if encontrado:
                        break
                
                if not encontrado:
                    if user.monto_credito % precio_adulto == 0:
                        user_dict['numero_personas'] = user.monto_credito // precio_adulto
                    elif user.monto_credito % precio_nino == 0:
                        user_dict['numero_personas'] = user.monto_credito // precio_nino
                    else:
                        mejor_diferencia = float('inf')
                        mejor_personas = 1
                        
                        for num_personas in range(1, 21):
                            for num_adultos in range(num_personas + 1):
                                num_ninos = num_personas - num_adultos
                                total = (num_adultos * precio_adulto) + (num_ninos * precio_nino)
                                diferencia = abs(total - user.monto_credito)
                                
                                if diferencia < mejor_diferencia:
                                    mejor_diferencia = diferencia
                                    mejor_personas = num_personas
                        
                        user_dict['numero_personas'] = mejor_personas

            elif user.tipo == 'asesoria+consular':
                # Este tipo tiene precio único por persona
                precio_persona = 1432900 - 100000  # 1,332,900
                
                if user.monto_credito % precio_persona == 0:
                    user_dict['numero_personas'] = user.monto_credito // precio_persona
                else:
                    # Buscar el número más cercano
                    user_dict['numero_personas'] = round(user.monto_credito / precio_persona)

            else:
                # Para otros tipos de crédito
                user_dict['numero_personas'] = None
                
            users_paginados.append(user_dict)

        # 4) Construcción de páginas visibles
        paginas_visibles = [
            p for p in range(max(1, pagina_actual - 2), min(total_paginas + 1, pagina_actual + 3))
        ]

        # 5) Renderizado
        return render_template(
            "list_credito.html",
            data=users_paginados,
            pagina_actual=pagina_actual,
            total_usuarios=total_usuarios,
            total_paginas=total_paginas,
            usuarios_por_pagina=usuarios_por_pagina,
            paginas_visibles=paginas_visibles,
        ), 200

    except Exception as e:
        app.logger.error("Error al obtener los créditos: %s", e)
        return render_template(
            "list_credito.html",
            data=[],
            pagina_actual=pagina_actual,
            total_usuarios=0,
            total_paginas=0,
            usuarios_por_pagina=usuarios_por_pagina,
            paginas_visibles=[],
        ), 200


@app.route("/list_all_credito", methods=["GET"])
@check_db_connection
@login_required
def list_all_credito():
    pagina_actual = request.args.get("pagina", default=1, type=int)
    usuarios_por_pagina = 10

    filtro_usuario = request.args.get("filtro_usuario", default="", type=str)
    filtro_tipo = request.args.get("filtro_tipo", default="", type=str)
    fecha_inicio = request.args.get("fecha_inicio", default="", type=str)
    fecha_fin = request.args.get("fecha_fin", default="", type=str)

    try:
        query_sql = db.session.query(
            Creditos.id_credito,
            Usuarios.nombre_completo,
            Creditos.monto_credito,
            Creditos.tasa_interes,
            Creditos.numero_cuotas,
            ModalidadesPago.nombre_modalidad,
            Creditos.fecha_credito,
            Creditos.fecha_creacion,
            Creditos.fecha_actualizacion,
            Creditos.tipo,
            func.count(
                case((Pagos.id_estado_deuda == 2, 1), else_=None)
            ).label('cuotas_pagadas'),
            exists().where(
                and_(
                    Archivos.id_usuario == Usuarios.id_usuario,
                    Archivos.id_subidor == 3
                )
            ).label('tiene_docs_id3')
        ).join(
            Usuarios, Creditos.id_usuario == Usuarios.id_usuario
        ).join(
            ModalidadesPago, Creditos.id_modalidad_pago == ModalidadesPago.id_modalidad
        ).outerjoin(
            Pagos, and_(
                Creditos.id_credito == Pagos.id_credito,
                Pagos.habilitado == True
            )
        ).filter(
            Creditos.habilitado == True
        )

        if filtro_usuario:
            query_sql = query_sql.filter(Usuarios.nombre_completo.ilike(f"%{filtro_usuario}%"))
        if filtro_tipo:
            query_sql = query_sql.filter(Creditos.tipo == filtro_tipo)
        if fecha_inicio:
            query_sql = query_sql.filter(Creditos.fecha_credito >= fecha_inicio)
        if fecha_fin:
            query_sql = query_sql.filter(Creditos.fecha_credito <= fecha_fin)

        query_sql = query_sql.group_by(
            Creditos.id_credito,
            Usuarios.nombre_completo,
            Creditos.monto_credito,
            Creditos.tasa_interes,
            Creditos.numero_cuotas,
            ModalidadesPago.nombre_modalidad,
            Creditos.fecha_credito,
            Creditos.fecha_creacion,
            Creditos.fecha_actualizacion,
            Creditos.tipo,
            Usuarios.id_usuario
        ).order_by(Creditos.id_credito.asc())

        total_usuarios = query_sql.count()
        total_paginas = (total_usuarios - 1) // usuarios_por_pagina + 1
        inicio = (pagina_actual - 1) * usuarios_por_pagina

        users_raw = query_sql.offset(inicio).limit(usuarios_por_pagina).all()

        users_paginados = []
        for user in users_raw:
            user_dict = user._asdict() if hasattr(user, '_asdict') else dict(user._mapping)

            tipo = user_dict["tipo"]
            monto = user_dict["monto_credito"]

            if tipo == 'normal':
                precio_adulto = 1096000 - 100000  # 996,000
                precio_nino = 694000 - 100000     # 594,000
                
                # Intentar encontrar la combinación exacta de adultos y niños
                encontrado = False
                
                # Probar diferentes combinaciones (hasta 5 personas total)
                for num_adultos in range(6):
                    for num_ninos in range(6 - num_adultos):
                        if num_adultos + num_ninos > 0:  # Al menos una persona
                            total_calculado = (num_adultos * precio_adulto) + (num_ninos * precio_nino)
                            if total_calculado == monto:
                                user_dict["numero_personas"] = num_adultos + num_ninos
                                encontrado = True
                                break
                    if encontrado:
                        break
                
                # Si no encontramos combinación exacta, aproximar
                if not encontrado:
                    # Primero intentar solo adultos
                    if monto % precio_adulto == 0:
                        user_dict["numero_personas"] = monto // precio_adulto
                    # Luego intentar solo niños
                    elif monto % precio_nino == 0:
                        user_dict["numero_personas"] = monto // precio_nino
                    else:
                        # Si no es divisible, buscar la mejor aproximación
                        mejor_diferencia = float('inf')
                        mejor_personas = 1
                        
                        for num_personas in range(1, 6):  # Máximo 5 personas
                            # Probar todas las combinaciones para este número de personas
                            for num_adultos in range(num_personas + 1):
                                num_ninos = num_personas - num_adultos
                                total = (num_adultos * precio_adulto) + (num_ninos * precio_nino)
                                diferencia = abs(total - monto)
                                
                                if diferencia < mejor_diferencia:
                                    mejor_diferencia = diferencia
                                    mejor_personas = num_personas
                        
                        user_dict["numero_personas"] = mejor_personas

            elif tipo == 'consular':
                precio_adulto = 1928500 - 385000  # 1,543,500
                precio_nino = 1526500 - 385000    # 1,141,500
                
                # Mismo enfoque: buscar combinación exacta primero
                encontrado = False
                
                for num_adultos in range(6):
                    for num_ninos in range(6 - num_adultos):
                        if num_adultos + num_ninos > 0:
                            total_calculado = (num_adultos * precio_adulto) + (num_ninos * precio_nino)
                            if total_calculado == monto:
                                user_dict["numero_personas"] = num_adultos + num_ninos
                                encontrado = True
                                break
                    if encontrado:
                        break
                
                if not encontrado:
                    if monto % precio_adulto == 0:
                        user_dict["numero_personas"] = monto // precio_adulto
                    elif monto % precio_nino == 0:
                        user_dict["numero_personas"] = monto // precio_nino
                    else:
                        mejor_diferencia = float('inf')
                        mejor_personas = 1
                        
                        for num_personas in range(1, 6):  # Máximo 5 personas
                            for num_adultos in range(num_personas + 1):
                                num_ninos = num_personas - num_adultos
                                total = (num_adultos * precio_adulto) + (num_ninos * precio_nino)
                                diferencia = abs(total - monto)
                                
                                if diferencia < mejor_diferencia:
                                    mejor_diferencia = diferencia
                                    mejor_personas = num_personas
                        
                        user_dict["numero_personas"] = mejor_personas

            elif tipo == 'asesoria':
                precio_adulto = 652000 - 100000  # 552,000
                precio_nino = 388000 - 100000    # 288,000
                
                # Mismo enfoque
                encontrado = False
                
                for num_adultos in range(6):
                    for num_ninos in range(6 - num_adultos):
                        if num_adultos + num_ninos > 0:
                            total_calculado = (num_adultos * precio_adulto) + (num_ninos * precio_nino)
                            if total_calculado == monto:
                                user_dict["numero_personas"] = num_adultos + num_ninos
                                encontrado = True
                                break
                    if encontrado:
                        break
                
                if not encontrado:
                    if monto % precio_adulto == 0:
                        user_dict["numero_personas"] = monto // precio_adulto
                    elif monto % precio_nino == 0:
                        user_dict["numero_personas"] = monto // precio_nino
                    else:
                        mejor_diferencia = float('inf')
                        mejor_personas = 1
                        
                        for num_personas in range(1, 6):  # Máximo 5 personas
                            for num_adultos in range(num_personas + 1):
                                num_ninos = num_personas - num_adultos
                                total = (num_adultos * precio_adulto) + (num_ninos * precio_nino)
                                diferencia = abs(total - monto)
                                
                                if diferencia < mejor_diferencia:
                                    mejor_diferencia = diferencia
                                    mejor_personas = num_personas
                        
                        user_dict["numero_personas"] = mejor_personas

            elif tipo == 'asesoria+consular':
                # Este tipo tiene precio único por persona
                precio_persona = 1432900 - 100000  # 1,332,900
                
                if monto % precio_persona == 0:
                    user_dict["numero_personas"] = monto // precio_persona
                else:
                    # Buscar el número más cercano
                    user_dict["numero_personas"] = round(monto / precio_persona)

            else:
                # Para otros tipos de crédito
                user_dict["numero_personas"] = None

            users_paginados.append(user_dict)

        # Calcular suma del monto total sin paginación
        query_suma = db.session.query(func.sum(Creditos.monto_credito)).join(
            Usuarios, Creditos.id_usuario == Usuarios.id_usuario
        ).filter(
            Creditos.habilitado == True
        )

        if filtro_usuario:
            query_suma = query_suma.filter(Usuarios.nombre_completo.ilike(f"%{filtro_usuario}%"))
        if filtro_tipo:
            query_suma = query_suma.filter(Creditos.tipo == filtro_tipo)
        if fecha_inicio:
            query_suma = query_suma.filter(Creditos.fecha_credito >= fecha_inicio)
        if fecha_fin:
            query_suma = query_suma.filter(Creditos.fecha_credito <= fecha_fin)

        suma_monto = query_suma.scalar() or 0

        paginas_visibles = [
            p for p in range(max(1, pagina_actual - 2), min(total_paginas + 1, pagina_actual + 3))
        ]

        return render_template(
            "list_all_credito.html",
            data=users_paginados,
            pagina_actual=pagina_actual,
            total_usuarios=total_usuarios,
            total_paginas=total_paginas,
            usuarios_por_pagina=usuarios_por_pagina,
            fecha_inicio=fecha_inicio,
            fecha_fin=fecha_fin,
            filtro_usuario=filtro_usuario,
            filtro_tipo=filtro_tipo,
            paginas_visibles=paginas_visibles,
            suma_monto=suma_monto
        ), 200

    except Exception as e:
        app.logger.error("Error al obtener todos los créditos: %s", e)
        return render_template(
            "list_all_credito.html",
            data=[],
            pagina_actual=pagina_actual,
            total_usuarios=0,
            total_paginas=0,
            usuarios_por_pagina=usuarios_por_pagina,
            fecha_inicio=fecha_inicio,
            fecha_fin=fecha_fin,
            filtro_usuario=filtro_usuario,
            filtro_tipo=filtro_tipo,
            paginas_visibles=[],
            suma_monto=0
        ), 200



@app.route('/download_creditos', methods=['GET'])
@check_db_connection
@login_required
def download_creditos():
    fecha_inicio = request.args.get("fecha_inicio", "")
    fecha_fin = request.args.get("fecha_fin", "")
    filtro_usuario = request.args.get("filtro_usuario", "").strip()

    try:
        # Realizamos los JOIN para incluir la información del usuario y de la modalidad de pago
        query = Creditos.query.join(Creditos.usuario).join(Creditos.modalidad_pago)

        # Filtro por rango de fechas en Creditos.fecha_credito
        if fecha_inicio and fecha_fin:
            fecha_inicio_dt = datetime.strptime(fecha_inicio, "%Y-%m-%d")
            fecha_fin_dt = datetime.strptime(fecha_fin, "%Y-%m-%d")
            query = query.filter(Creditos.fecha_credito.between(fecha_inicio_dt, fecha_fin_dt))

        # Filtro por nombre de usuario (campo de la tabla Usuarios)
        if filtro_usuario:
            query = query.filter(Usuarios.nombre_completo.ilike(f"%{filtro_usuario}%"))

        # Obtener los resultados
        data = query.all()

        # Crear un archivo Excel en memoria
        output = io.BytesIO()
        workbook = xlsxwriter.Workbook(output, {'in_memory': True})
        worksheet = workbook.add_worksheet()

        # Formatos para la cabecera y fechas
        header_format = workbook.add_format({'bold': True, 'align': 'center', 'border': 1})
        date_format = workbook.add_format({'num_format': 'yyyy-mm-dd'})
        datetime_format = workbook.add_format({'num_format': 'yyyy-mm-dd hh:mm:ss'})

        # Encabezados del Excel
        headers = [
            'ID', 'USUARIO', 'CREDITO', 'INTERES', 'CUOTAS', 'MODALIDAD',
            'GANANCIA', 'FECHA CREDITO', 'FECHA CREACION', 'FECHA ACTUALIZACION', 'HABILITADO'
        ]
        for col_num, header in enumerate(headers):
            worksheet.write(0, col_num, header, header_format)

        # Preparar los datos para escribirlos en el Excel
        row_data = []
        for credito in data:
            row = [
                credito.id_credito,
                credito.usuario.nombre_completo,        # Obtenemos el nombre desde Usuarios
                credito.monto_credito,
                credito.tasa_interes,
                credito.numero_cuotas,
                credito.modalidad_pago.nombre_modalidad,  # Obtenemos el nombre de la modalidad
                credito.ganancia_total,
                credito.fecha_credito,
                credito.fecha_creacion,
                credito.fecha_actualizacion,
                'Sí' if credito.habilitado else 'No'
            ]
            row_data.append(row)

        # Escribir cada fila en el Excel
        for row_num, row in enumerate(row_data, start=1):
            for col_num, cell_data in enumerate(row):
                if isinstance(cell_data, datetime):
                    # Seleccionamos el formato adecuado según si tiene hora o no
                    if cell_data.hour == 0 and cell_data.minute == 0 and cell_data.second == 0:
                        worksheet.write_datetime(row_num, col_num, cell_data, date_format)
                    else:
                        worksheet.write_datetime(row_num, col_num, cell_data, datetime_format)
                elif isinstance(cell_data, date) and not isinstance(cell_data, datetime):
                    worksheet.write_datetime(row_num, col_num, datetime.combine(cell_data, datetime.min.time()), date_format)
                elif cell_data is None:
                    worksheet.write(row_num, col_num, '')
                else:
                    worksheet.write(row_num, col_num, cell_data)

        # Ajustar automáticamente el ancho de las columnas según el contenido
        for col_num in range(len(headers)):
            max_length = max(len(str(row[col_num])) for row in row_data) + 5
            worksheet.set_column(col_num, col_num, max_length)

        # Cerrar el libro y preparar la respuesta
        workbook.close()
        output.seek(0)
        return Response(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition": "attachment;filename=creditos.xlsx"}
        )
    except Exception as e:
        print("Error al generar el archivo Excel:", e)
        return "Error al generar el archivo Excel", 500


@app.route("/list_all_pagos", methods=["GET"])
@check_db_connection
@login_required
def list_all_pagos():
    pagina_actual = request.args.get("pagina", default=1, type=int)
    usuarios_por_pagina = 10

    # Filtros opcionales
    filtro_usuario = request.args.get("filtro_usuario", "").strip()
    filtro_estado = request.args.get("filtro_estado", "").strip()
    fecha_inicio = request.args.get("fecha_inicio", "").strip()
    fecha_fin = request.args.get("fecha_fin", "").strip()

    try:
        query = db.session.query(
            Pagos.id_pago,
            Usuarios.nombre_completo,
            Pagos.numero_cuota,
            Pagos.fecha_pago_programada,
            Pagos.fecha_pago_real,
            Pagos.monto_cuota,
            Pagos.abono_capital,
            Pagos.abono_intereses,
            Pagos.saldo_restante,
            EstadosDeuda.nombre_estado,
            Usuarios.id_usuario,
            Creditos.id_credito,
            Pagos.dias_mora
        ).select_from(Pagos) \
        .join(Creditos, Pagos.id_credito == Creditos.id_credito) \
        .join(Usuarios, Creditos.id_usuario == Usuarios.id_usuario) \
        .join(EstadosDeuda, Pagos.id_estado_deuda == EstadosDeuda.id_estado) \
        .filter(Pagos.habilitado == True)

        # Aplicar filtros dinámicos
        if filtro_usuario:
            query = query.filter(
                or_(
                    Usuarios.nombre_completo.ilike(f"%{filtro_usuario}%"),
                    Usuarios.cedula.ilike(f"%{filtro_usuario}%"),
                    Usuarios.celular.ilike(f"%{filtro_usuario}%")
                )
            )
        if filtro_estado:
            query = query.filter(EstadosDeuda.nombre_estado.ilike(f"%{filtro_estado}%"))
        if fecha_inicio and fecha_fin:
            query = query.filter(Pagos.fecha_pago_programada.between(fecha_inicio, fecha_fin))
    
         # Calcular la suma total del monto de la cuota según el filtro
        suma_monto = query.with_entities(db.func.sum(Pagos.monto_cuota)).scalar() or 0

        total_usuarios = query.count()
        total_paginas = (total_usuarios - 1) // usuarios_por_pagina + 1
        inicio = (pagina_actual - 1) * usuarios_por_pagina
        users_paginados = query.offset(inicio).limit(usuarios_por_pagina).all()

        paginas_visibles = [
            pagina for pagina in range(
                max(1, pagina_actual - 2),
                min(total_paginas + 1, pagina_actual + 3)
            )
        ]

        return render_template(
            "list_all_pagos.html",
            data=users_paginados,
            pagina_actual=pagina_actual,
            total_usuarios=total_usuarios,
            total_paginas=total_paginas,
            usuarios_por_pagina=usuarios_por_pagina,
            filtro_usuario=filtro_usuario,
            filtro_estado=filtro_estado,
            fecha_inicio=fecha_inicio, 
            fecha_fin=fecha_fin,
            paginas_visibles=paginas_visibles,
            suma_monto=suma_monto  # Pasar la suma a la plantilla

        ), 200

    except Exception as e:
        print("Error al obtener los pagos:", e)
        return render_template(
            "list_all_pagos.html",
            data=[],
            pagina_actual=pagina_actual,
            total_usuarios=0,
            total_paginas=0,
            usuarios_por_pagina=usuarios_por_pagina,
            filtro_usuario=filtro_usuario,
            filtro_estado=filtro_estado,
            fecha_inicio=fecha_inicio, 
            fecha_fin=fecha_fin,
            paginas_visibles=[],
            suma_monto=suma_monto  # Pasar la suma a la plantilla

        ), 200

@app.route('/download_usuarios', methods=['GET'])
@check_db_connection
@login_required
def download_usuarios():
    try:
        # Obtener todos los usuarios
        usuarios = Usuarios.query.all()

        # Crear archivo Excel en memoria
        output = io.BytesIO()
        workbook = xlsxwriter.Workbook(output, {'in_memory': True})
        worksheet = workbook.add_worksheet("Usuarios")

        # Formato para cabecera y fechas
        header_format = workbook.add_format({'bold': True, 'align': 'center', 'border': 1})
        datetime_format = workbook.add_format({'num_format': 'yyyy-mm-dd hh:mm:ss'})

        # Encabezados
        headers = [
            'ID', 'Nombre completo', 'Cédula', 'Dirección', 'Celular', 'Correo',
            'Expedición', 'Residencia', 'Habilitado', 'Fecha creación', 'Fecha actualización'
        ]

        for col_num, header in enumerate(headers):
            worksheet.write(0, col_num, header, header_format)

        # Escribir datos
        for row_num, usuario in enumerate(usuarios, start=1):
            row = [
                usuario.id_usuario,
                usuario.nombre_completo,
                usuario.cedula,
                usuario.direccion,
                usuario.celular,
                usuario.correo,
                usuario.expedicion,
                usuario.residencia,
                'Sí' if usuario.habilitado else 'No',
                usuario.fecha_creacion,
                usuario.fecha_actualizacion
            ]
            for col_num, cell_data in enumerate(row):
                if isinstance(cell_data, datetime):
                    worksheet.write_datetime(row_num, col_num, cell_data, datetime_format)
                elif cell_data is None:
                    worksheet.write(row_num, col_num, '')
                else:
                    worksheet.write(row_num, col_num, cell_data)

        # Ajustar anchos de columna
        for col_num in range(len(headers)):
            worksheet.set_column(col_num, col_num, 20)

        workbook.close()
        output.seek(0)

        return Response(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition": "attachment;filename=usuarios.xlsx"}
        )

    except Exception as e:
        print("Error al generar el archivo Excel de usuarios:", e)
        return "Error al generar el archivo Excel de usuarios", 500

@app.route('/download_pagos', methods=['GET'])
@check_db_connection
@login_required
def download_pagos():
    try:
        
        query = Pagos.query.join(Pagos.credito).join(Creditos.usuario).join(Pagos.estado_deuda)
        
        # Si deseas aplicar filtros, por ejemplo por fecha programada, podrías hacerlo así:
        filtro_usuario = request.args.get("filtro_usuario", "").strip()
        fecha_inicio = request.args.get("fecha_inicio", "")
        fecha_fin = request.args.get("fecha_fin", "")
        if fecha_inicio and fecha_fin:
            fecha_inicio_dt = datetime.strptime(fecha_inicio, "%Y-%m-%d").date()
            fecha_fin_dt = datetime.strptime(fecha_fin, "%Y-%m-%d").date()
            query = query.filter(Pagos.fecha_pago_programada.between(fecha_inicio_dt, fecha_fin_dt))
        if filtro_usuario:
            query = query.filter(Usuarios.nombre_completo.ilike(f"%{filtro_usuario}%"))
        
        # Obtener todos los registros
        data = query.all()
        
        # Crear un archivo Excel en memoria
        output = io.BytesIO()
        workbook = xlsxwriter.Workbook(output, {'in_memory': True})
        worksheet = workbook.add_worksheet()
        
        # Formatos para cabecera y fechas
        header_format = workbook.add_format({'bold': True, 'align': 'center', 'border': 1})
        date_format = workbook.add_format({'num_format': 'yyyy-mm-dd'})
        datetime_format = workbook.add_format({'num_format': 'yyyy-mm-dd hh:mm:ss'})
        
        # Definir los encabezados de las columnas
        headers = [
            'ID', 'Usuario', 'Cuota', 'Fecha programada', 'Fecha pago',
            'Monto', 'Abono capital', 'Abono intereses', 'Estado', 'Días mora'
        ]
        for col_num, header in enumerate(headers):
            worksheet.write(0, col_num, header, header_format)
        
        # Preparar los datos para cada fila
        row_data = []
        for pago in data:
            row = [
                pago.id_pago,
                pago.credito.usuario.nombre_completo,    # Obtiene el nombre del usuario
                pago.numero_cuota,
                pago.fecha_pago_programada,               # Tipo date
                pago.fecha_pago_real,                     # Tipo date (o None)
                pago.monto_cuota,
                pago.abono_capital,
                pago.abono_intereses,
                pago.estado_deuda.nombre_estado,          # Estado del pago (Pagado, Pendiente, Vencido)
                pago.dias_mora
            ]
            row_data.append(row)
        
        # Escribir los datos en la hoja de cálculo
        for row_num, row in enumerate(row_data, start=1):
            for col_num, cell_data in enumerate(row):
                # Si se trata de un objeto datetime o date se formatea adecuadamente
                if isinstance(cell_data, datetime):
                    if cell_data.hour == 0 and cell_data.minute == 0 and cell_data.second == 0:
                        worksheet.write_datetime(row_num, col_num, cell_data, date_format)
                    else:
                        worksheet.write_datetime(row_num, col_num, cell_data, datetime_format)
                elif isinstance(cell_data, date) and not isinstance(cell_data, datetime):
                    worksheet.write_datetime(row_num, col_num, datetime.combine(cell_data, datetime.min.time()), date_format)
                elif cell_data is None:
                    worksheet.write(row_num, col_num, '')
                else:
                    worksheet.write(row_num, col_num, cell_data)
        
        # Ajustar el ancho de las columnas según el contenido
        for col_num in range(len(headers)):
            max_length = max(len(str(row[col_num])) for row in row_data) + 5
            worksheet.set_column(col_num, col_num, max_length)
        
        # Finalizar y preparar el archivo para enviar
        workbook.close()
        output.seek(0)
        
        return Response(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={"Content-Disposition": "attachment;filename=pagos.xlsx"}
        )
    except Exception as e:
        print("Error al generar el archivo Excel de pagos:", e)
        return "Error al generar el archivo Excel de pagos", 500


@app.route("/history_user/<int:id>", methods=["GET", "POST"])
@check_db_connection
@login_required
def history_user(id):
    
    usuario_actual = Usuarios.query.get(id)
    if not usuario_actual:
        return "Usuario no encontrado", 404
    
    pagina_actual = request.args.get("pagina", default=1, type=int)
    usuarios_por_pagina = 10
    
    try:
        query = db.session.query(
            Auditoria,
            Administradores.correo
        ).join(
            Administradores, Administradores.id_administrador == Auditoria.id_administrador
        ).filter(
            Auditoria.id_usuario == id
        )
        
        total_usuarios = query.count()
        total_paginas = (total_usuarios - 1) // usuarios_por_pagina + 1
        inicio = (pagina_actual - 1) * usuarios_por_pagina
        users_paginados = query.offset(inicio).limit(usuarios_por_pagina).all()

        paginas_visibles = [
            pagina for pagina in range(
                max(1, pagina_actual - 2),
                min(total_paginas + 1, pagina_actual + 3)
            )
        ]


        return render_template('history_user.html', histories=users_paginados, pagina_actual=pagina_actual, paginas_visibles=paginas_visibles, total_paginas=total_paginas), 200
    except Exception as e:
        print("Error al obtener el historial:", e)
        return render_template('history_user.html', histories=[],pagina_actual=pagina_actual, paginas_visibles=[], total_paginas=0), 200


@app.route("/update_user/<int:id>", methods=["GET", "POST"])
@check_db_connection
@login_required
def update_user(id):
    mostrar_exito = False
    mostrar_alerta = False
    
    usuario_actual = Usuarios.query.get(id)

    if not usuario_actual:
        return "Usuario no encontrado", 404
    
    if request.method == "POST":
        # Obtener valores actuales y nuevos
        cambios = {}
        nuevos_valores = {
            "nombre_completo": request.form.get("nombre_completo"),
            "cedula": request.form.get("cedula"),
            "direccion": request.form.get("direccion"),
            "celular": request.form.get("celular"),
            "correo": request.form.get("correo"),
        }

        # Comparar valores actuales con los nuevos
        for campo, nuevo_valor in nuevos_valores.items():
            valor_actual = getattr(usuario_actual, campo)
            if str(valor_actual) != str(nuevo_valor):  # Compara valores como cadenas
                cambios[campo] = nuevo_valor
                setattr(usuario_actual, campo, nuevo_valor)  # Actualiza solo si hay cambios
                
        
        # Si hay cambios, actualiza la base de datos y registra en la auditoría
        if cambios:
            try:
                db.session.commit()
                # Registrar en la auditoría
                auditoria = Auditoria(
                    id_administrador=current_user.id_administrador,  # ID del administrador actual
                    tabla_afectada="Usuarios",
                    id_usuario=id,
                    cambios=cambios,
                    tipo_operacion="Actualizar"
                )
                db.session.add(auditoria)
                db.session.commit()
                mostrar_exito = True
            except Exception:
                db.session.rollback()
                mostrar_alerta = True
        
        # ----- Procesar subida de archivos -----
        files = request.files.getlist('file')

        if files:
            try:
                for i, file in enumerate(files):
                    if file and file.filename:
                        # Obtener la fecha correspondiente a este archivo
                        upload_date_str = request.form.get(f'upload_date_{i}')
                        if not upload_date_str:
                            raise ValueError(f"Falta la fecha de subida para el archivo {file.filename}")
                        
                        fecha_subida = datetime.strptime(upload_date_str, '%Y-%m-%d').date()

                        nombre_seguro = secure_filename(file.filename)
                        unique_name = f"{id}_{nombre_seguro}"
                        ruta = os.path.join('pdfs', unique_name)
                        file.save(ruta)

                        # Crear registro en BD
                        nuevo = Archivos(
                            id_usuario=id,
                            nombre_archivo=unique_name,
                            id_subidor=current_user.id_administrador,
                            correo_subidor=current_user.correo,
                            fecha_subida=fecha_subida
                        )
                        db.session.add(nuevo)

                db.session.commit()
                mostrar_exito = True
                
                

                if current_user.id_administrador == 3:
                    # Aquí buscamos el crédito más reciente de este usuario
                    ultimo_credito = (
                        Creditos.query
                        .filter_by(id_usuario=id)
                        .order_by(Creditos.fecha_credito.desc())
                        .first()
                    )
                    message = f"Pago de crédito registrado con éxito - {usuario_actual.nombre_completo}"
                    target_user_id = 1
                    send_notification(target_user_id, message, credit_id=ultimo_credito.id_credito)

            except Exception as e:
                db.session.rollback()
                print("Error al subir archivos:", e)
                mostrar_alerta = True

        else:
            mostrar_alerta = True
            
        return render_template('update_user.html', user=usuario_actual, mostrar_exito=mostrar_exito, mostrar_alerta=mostrar_alerta)
    
    elif request.method == "GET":
        return render_template('update_user.html', user=usuario_actual)



    

@app.route("/deshabilitar_usuario/<int:id>", methods=["PUT"])
@check_db_connection
@login_required
def deshabilitar_usuario(id):
    if current_user.id_administrador != 1 and current_user.id_administrador !=2:
        return "No tienes permiso", 403
    if request.method == "PUT":
        try:
            usuario = Usuarios.query.get(id)
            if usuario:
                usuario.habilitado = False
                db.session.commit()
                
                auditoria = Auditoria(
                    id_administrador=current_user.id_administrador,  # ID del administrador actual
                    tabla_afectada="Usuarios",
                    id_usuario=id,
                    cambios= {"Habilitado": "No"},
                    tipo_operacion="Eliminar"
                )
                db.session.add(auditoria)
                db.session.commit()
                
                return "", 204  # No Content
            else:
                return "Usuario no encontrado", 404  # Not Found
        except Exception as error:
            print("Error al deshabilitar el usuario:", error)
            db.session.rollback()
            return "Error al deshabilitar el usuario", 500  # Internal Server Error
    else:
        return "Método no permitido", 405  # Method Not Allowed


@app.route("/deshabilitar_credito/<int:id>", methods=["PUT"])
@check_db_connection
@login_required
def deshabilitar_credito(id):
    if current_user.id_administrador != 1 and current_user.id_administrador !=2:
        return "No tienes permiso", 403
    if request.method == "PUT":
        try:
            credito = Creditos.query.get(id)
            

            if credito:
                credito.habilitado = False
                db.session.commit()
                
                auditoria = Auditoria(
                    id_administrador=current_user.id_administrador,  # ID del administrador actual
                    tabla_afectada="Creditos",
                    id_usuario=credito.id_usuario,
                    cambios= {"Habilitado": "No",
                              "id credito": credito.id_credito
                            },
                    tipo_operacion="Eliminar"
                )
                db.session.add(auditoria)
                db.session.commit()
                
                return "", 204  # No Content
            else:
                return "Credito no encontrado", 404  # Not Found
        except Exception as error:
            print("Error al deshabilitar el credito:", error)
            db.session.rollback()
            return "Error al deshabilitar el credito", 500  # Internal Server Error
    else:
        return "Método no permitido", 405  # Method Not Allowed
    
    
@app.route('/usuario/<int:id_usuario>/archivo', methods=['POST'])
def subir_archivos(id_usuario):
    if 'file' not in request.files:
        return jsonify({'message': 'No se envió ningún archivo'}), 400
    
    files = request.files.getlist('file')
    if not files:
        return jsonify({'message': 'No se enviaron archivos válidos'}), 400

    saved_files = []
    for file in files:
        if file.filename == '':
            continue  # Omite archivos sin nombre
        if file and file.filename.lower().endswith('.pdf'):
            filename = secure_filename(file.filename)
            # Genera un nombre único (por ejemplo, anteponiendo el id_usuario)
            unique_filename = f"{id_usuario}_{filename}"
            filepath = os.path.join(UPLOAD_FOLDER, unique_filename)
            file.save(filepath)
            saved_files.append(unique_filename)
        else:
            return jsonify({'message': 'Formato de archivo no permitido. Solo se aceptan PDFs.'}), 400

    if saved_files:
        return jsonify({'message': 'Archivos guardados correctamente', 'archivos': saved_files}), 200
    else:
        return jsonify({'message': 'Ningún archivo válido fue subido'}), 400


@app.route('/usuario/<int:id_usuario>/archivos', methods=['GET'])
@login_required
def listar_archivos(id_usuario):
    try:
        archivos_bd = Archivos.query.filter_by(id_usuario=id_usuario) \
                                   .order_by(Archivos.fecha_subida.desc()).all()
                                   
        resultado = []

            
        for f in archivos_bd:
            resultado.append({
                'filename': f.nombre_archivo,
                'uploader_email': f.correo_subidor,
                'upload_date': f.fecha_subida.strftime('%d/%m/%Y'),
                'creation_date': f.fecha_creacion.strftime('%d/%m/%Y %H:%M:%S'),
                'download_url': url_for('descargar_archivo',
                                        id_usuario=id_usuario,
                                        filename=f.nombre_archivo,
                                        _external=False)
            })
            print(resultado)
        return jsonify(resultado), 200
    except Exception as e:
        return jsonify({'message': 'Error al listar archivos', 'error': str(e)}), 500



@app.route('/usuario/<int:id_usuario>/archivo/<filename>', methods=['GET'])
@login_required
def descargar_archivo(id_usuario, filename):
    # Validar metadata en BD
    archivo = Archivos.query.filter_by(
        id_usuario=id_usuario,
        nombre_archivo=filename
    ).first()
    if not archivo:
        return jsonify({'message': 'Archivo no encontrado o sin permiso'}), 404

    carpeta = app.config['UPLOAD_FOLDER']
    if os.path.exists(os.path.join(carpeta, filename)):
        return send_from_directory(carpeta, filename, as_attachment=True, download_name=filename)
    else:
        return jsonify({'message': 'El archivo no existe en el servidor'}), 500



@app.route('/usuario/<int:id_usuario>/archivo/<string:filename>', methods=['DELETE'])
@login_required
def eliminar_archivo(id_usuario, filename):
    if not filename.startswith(f"{id_usuario}_"):
        return jsonify({'message': 'No autorizado para eliminar este archivo.'}), 403

    filepath = os.path.join(UPLOAD_FOLDER, filename)

    if os.path.exists(filepath):
        try:
            os.remove(filepath)

            # Eliminar el registro en la base de datos
            archivo = Archivos.query.filter_by(id_usuario=id_usuario, nombre_archivo=filename).first()
            if archivo:
                db.session.delete(archivo)
                db.session.commit()

            return jsonify({'message': 'Archivo eliminado correctamente'}), 200
        except Exception as e:
            return jsonify({'message': 'Error al eliminar el archivo', 'error': str(e)}), 500
    else:
        return jsonify({'message': 'Archivo no encontrado'}), 404




def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


@app.route('/pago/<int:id_pago>/archivo', methods=['GET', 'POST'])
def gestionar_pago(id_pago):
    
    # Ruta completa del archivo
    filepath = os.path.join(UPLOAD_FOLDER_PAGOS, f"{id_pago}.jpg")

    if request.method == 'GET':
        # Descargar el archivo si existe
        if os.path.exists(filepath):
            return send_file(filepath, as_attachment=True, download_name=f"{id_pago}.jpg")
        else:
            # Responder con 404 si no se encuentra el archivo
            return jsonify({'message': 'Archivo no encontrado'}), 404

    if request.method == 'POST':
        # Subir o reemplazar el archivo
        if 'file' not in request.files:
            return jsonify({'message': 'No se envió ningún archivo'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'message': 'Nombre de archivo inválido'}), 400

        if file.filename != '' and allowed_file(file.filename):
             # Convertir la imagen a un formato estándar (JPEG)
            image = Image.open(file)
            filepath = os.path.join(UPLOAD_FOLDER_PAGOS, f"{id}.jpg")  # Guardar como JPG
            
            # Convertir y guardar la imagen en JPEG
            image.convert('RGB').save(filepath, "JPEG")
            return jsonify({'message': 'Archivo guardado correctamente'}), 200
        else:
            return jsonify({'message': 'Formato de archivo no permitido. Solo se aceptan imagenes.'}), 400
        
        
# Función para enviar notificaciones
def send_notification(target_user_id, message, credit_id=None):
    if str(target_user_id) in current_active_users:
        socketio.emit('notification', {'message': message, 'credit_id': credit_id}, room=str(target_user_id))
        
    notification = Notification(
        user_id=target_user_id,
        message=message,
        credit_id=credit_id,
        status='pendiente'
    )
    db.session.add(notification)
    db.session.commit()


@app.route("/notificaciones")
@login_required
def notifications():
    pagina_actual = request.args.get("pagina", default=1, type=int)
    notificaciones_por_pagina = 10
    admin_id = current_user.id_administrador

    # 1) Trae TODAS las notificaciones de este admin, ordenadas de más recientes a más viejas
    todas = (
        Notification.query
        .filter_by(user_id=admin_id)
        .order_by(Notification.fecha_creacion.desc())
        .all()
    )

    # 2) Construye un dict para quedarte sólo la PRIMERA notificación de cada credit_id
    unique_per_credit = {}
    for notif in todas:
        cid = notif.credit_id
        # clave única: credit_id o el propio id de notif si no tiene crédito
        key = cid if cid is not None else f"none_{notif.id_notificaciones}"
        if key not in unique_per_credit:
            unique_per_credit[key] = notif

    filtradas = list(unique_per_credit.values())
    # 3) Si es admin=3, quitamos aquí los créditos que ya tienen archivos subidos por el admin 3
    if admin_id == 3:
        filtradas = [
            notif for notif in filtradas
            if not (
                notif.credit_id and
                # buscamos si para ese crédito existe un archivo subido por admin=3
                db.session.query(Archivos)
                  .join(Creditos, Creditos.id_usuario == Archivos.id_usuario)
                  .filter(
                      Creditos.id_credito == notif.credit_id,
                      Archivos.id_subidor == 3
                  )
                  .first()
            )
        ]

    # 4) Ordenamos de nuevo por fecha
    filtradas.sort(key=lambda n: n.fecha_creacion, reverse=True)

    # 5) Paginación “manual”
    total_notifs = len(filtradas)
    total_paginas = (total_notifs - 1) // notificaciones_por_pagina + 1
    inicio = (pagina_actual - 1) * notificaciones_por_pagina
    fin = inicio + notificaciones_por_pagina
    notifs_paginadas = filtradas[inicio:fin]

    # 6) Para cada notif con credit_id, carga usuario_id
    for notif in notifs_paginadas:
        if notif.credit_id:
            credito = Creditos.query.get(notif.credit_id)
            notif.usuario_id = credito.id_usuario if credito else None
        else:
            notif.usuario_id = None

    # 7) Marca todas las pendientes como leídas
    Notification.query.filter_by(
        user_id=admin_id, status='pendiente'
    ).update({'status': 'leido'})
    db.session.commit()

    # 8) Renderizado
    paginas_visibles = list(range(
        max(1, pagina_actual - 2),
        min(total_paginas + 1, pagina_actual + 3)
    ))
    return render_template(
        "notificaciones.html",
        notifications=notifs_paginadas,
        pagina_actual=pagina_actual,
        paginas_visibles=paginas_visibles,
        total_paginas=total_paginas
    ), 200




# Eventos de Socket.IO para la conexión y desconexión
@socketio.on('join')
def handle_join(user_id):
    from flask_socketio import join_room
    join_room(str(user_id))
    current_active_users.add(str(user_id))
    print(f"Usuario {user_id} se ha unido a su sala.")

@socketio.on('disconnect')
def handle_disconnect():
    # Aquí deberías remover al usuario de current_active_users si tienes la información.
    pass

@app.route("/api/notificaciones_pendientes")
@login_required
def api_notificaciones_pendientes():
        
    pending_count = Notification.query.filter_by(
        user_id=current_user.id_administrador,
        status='pendiente'
    ).count()
    return jsonify({'pending_count': pending_count})


def migrar_archivos():
    archivos = os.listdir(UPLOAD_FOLDER)
    total = 0

    for archivo in archivos:
        try:
            if '_' not in archivo:
                continue  # no sigue la convención de nombres

            partes = archivo.split('_', 1)
            id_usuario = int(partes[0])
            nombre_seguro = secure_filename(archivo)
            ruta_archivo = os.path.join(UPLOAD_FOLDER, archivo)

            # Fecha de creación del archivo en el sistema
            timestamp_creacion = os.path.getctime(ruta_archivo)
            fecha_creacion = datetime.fromtimestamp(timestamp_creacion)

            # Verifica si ya existe en la base de datos
            ya_existe = Archivos.query.filter_by(nombre_archivo=archivo).first()
            if ya_existe:
                continue  # evita duplicados

            nuevo = Archivos(
                id_usuario=id_usuario,
                nombre_archivo=archivo,
                fecha_creacion=fecha_creacion,
                fecha_subida=fecha_creacion,  # puedes cambiarlo si tienes mejor dato
                id_subidor=1,
                correo_subidor="Fernando@reagendar.com"
            )
            db.session.add(nuevo)
            total += 1

        except Exception as e:
            print(f"Error procesando {archivo}: {e}")

    db.session.commit()
    print(f"{total} archivos migrados exitosamente.")
    
@app.route('/migrar-archivos')
@login_required
def migrar_archivos_view():
    if current_user.id_administrador != 2:
        return "Acceso denegado", 403
    migrar_archivos()
    return "Archivos migrados correctamente"
